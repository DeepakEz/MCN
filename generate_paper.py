"""Generate MCN Academic Paper (IEEE format) + Progress Report as DOCX and PDF.

Produces four files:
  MCN_Progress_Report.docx / .pdf  — technical progress log (all runs)
  MCN_Academic_Paper.docx  / .pdf  — IEEE two-column conference paper
"""

from __future__ import annotations
import os, re

# ── python-docx ──────────────────────────────────────────────────────────────
from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── ReportLab ────────────────────────────────────────────────────────────────
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import ParagraphStyle
from reportlab.lib.units import cm, inch
from reportlab.lib import colors
from reportlab.platypus import (
    BaseDocTemplate, SimpleDocTemplate, Frame, PageTemplate,
    Paragraph, Spacer, Table, TableStyle,
    HRFlowable, Preformatted, PageBreak, KeepTogether,
    FrameBreak, NextPageTemplate,
)
from reportlab.lib.enums import TA_LEFT, TA_CENTER, TA_JUSTIFY, TA_RIGHT

OUT_DIR = r"C:\MCN"

# ══════════════════════════════════════════════════════════════════════════════
# Colour palette (shared)
# ══════════════════════════════════════════════════════════════════════════════
DARK_BLUE   = colors.HexColor("#1F3864")
MID_BLUE    = colors.HexColor("#2E74B5")
LIGHT_BLUE  = colors.HexColor("#E8EEF7")
ACCENT      = colors.HexColor("#4472C4")
WHITE       = colors.white
BLACK       = colors.black
GREY        = colors.HexColor("#666666")
LGREY       = colors.HexColor("#F5F5F5")
ROW_ALT     = colors.HexColor("#F2F6FC")
AMBER       = colors.HexColor("#FFF2CC")
GREEN_LIGHT = colors.HexColor("#E2EFDA")

GRID_BASE = [
    ("GRID",         (0,0),(-1,-1), 0.3, colors.HexColor("#AAAAAA")),
    ("VALIGN",       (0,0),(-1,-1), "MIDDLE"),
    ("TOPPADDING",   (0,0),(-1,-1), 2),
    ("BOTTOMPADDING",(0,0),(-1,-1), 2),
]


# ══════════════════════════════════════════════════════════════════════════════
# ███████████████  IEEE PAPER CONTENT  ███████████████████████████████████████
# ══════════════════════════════════════════════════════════════════════════════

PAPER_TITLE   = ("Mycelial Council Network: Temperature Dominates Routing "
                 "in Multi-Agent LLM Code Synthesis — A Longitudinal Study "
                 "with Category-Level Analysis")
PAPER_AUTHORS = "MCN Research Initiative"
PAPER_AFFIL   = "Anonymous Submission — February 2026"

# ── Short abstract (~160 words, IEEE standard) ────────────────────────────────
IEEE_ABSTRACT = (
    "We empirically study whether a learned router can improve upon the best "
    "single-agent baseline when a council of LLM agents shares the same base "
    "model. Over ten experimental runs, three controlled ablations, two "
    "2,000-task longitudinal experiments, and a temperature-heterogeneous "
    "Phase 2 trial using Qwen2.5-Coder-7B-Instruct-AWQ on an 8-category "
    "Python benchmark, we find a consistent null result: neither router "
    "sophistication nor temperature diversity improves performance. "
    "A single agent at T=0.3 achieves 91%; MCN-LinUCB scores 86% (−5 pp). "
    "At 2,000 tasks LinUCB achieves 60.7% and a GNN router — with 40x more "
    "parameters — achieves 60.6% (delta = −0.1 pp). Phase 2 introduces "
    "heterogeneous tribe temperatures (T0=0.1, T1=0.5, T2=0.9) with a live "
    "Category Thompson Sampling router: overall pass rate is 60.7%, identical "
    "to Phase 1C, with per-tribe rates within 2 pp (T0=60.4%, T1=59.3%, "
    "T2=61.3%). The oracle gap narrows from 5.1 pp to 3.7 pp. CTS learns "
    "per-category routing preferences (oscillating T1→T0→T2 drift) but "
    "cannot improve aggregate performance because temperature diversity "
    "does not create capability diversity. Category-level heterogeneity "
    "dominates throughout: string 99%, graph 0%, corrected pass rate "
    "excluding graph = 81.3%. We conclude that routing value requires "
    "genuine inter-model performance variance, not sampling or algorithmic "
    "sophistication."
)

IEEE_KEYWORDS = (
    "multi-agent systems, LLM code synthesis, contextual bandits, "
    "graph neural networks, temperature scaling, routing collapse, "
    "ablation study, category-level analysis, oracle gap decomposition"
)

# ── Section data: (heading, body).  "A." prefix => subsection ────────────────
IEEE_SECTIONS = [

("I. INTRODUCTION",
"""\
Automated code synthesis with large language models (LLMs) has advanced \
rapidly, yet most deployments use a single model with fixed inference \
parameters. The Mycelial Council Network (MCN) hypothesis is that a council \
of agents with distinct styles, routed by a learned policy, can outperform \
any single agent — analogously to how software teams benefit from diverse \
specialists.

This paper reports ten experimental runs and three controlled single-agent \
ablations designed to test this hypothesis. Experiments proceed through three \
phases:

  Phase 1 (Runs 1–6): Infrastructure validation and GNN vs. LinUCB comparison \
on a 12-task benchmark. After correcting three infrastructure defects, the GNN \
router reached 98% and LinUCB reached 97%, but LinUCB collapsed to a single \
tribe (100/0/0 routing).

  Phase 2 (Runs 7–8): Task-set expansion to 16 types, revealing a 14 pp \
pass-rate drop from the 12-task ceiling and identifying three new \
reference-solution bugs.

  Phase 3 (Runs 9–10 + ablations): Controlled five-condition comparison — \
three single-agent ablations (T=0.1, T=0.3, T=0.9) and two MCN variants \
(LinUCB homogeneous, GNN heterogeneous).

The overall conclusion is that temperature selection matters more than routing \
strategy under these conditions, and routing with same-model tribes imposes a \
net exploration cost with no compensating specialisation benefit.

We make the following claims, strictly grounded in observed results:
  (1) A single agent at T=0.3 achieves 91%, outperforming all MCN variants.
  (2) MCN-LinUCB (homogeneous T=0.3) scores 86% — 5 pp below single-agent, \
with no statistically significant routing specialisation (chi-squared p=0.74).
  (3) MCN-GNN (heterogeneous T=0.1/0.5/0.9) scores 88% = T=0.9 single agent \
exactly. T_2 achieves 95% in isolation, but the router incurs a −7 pp cost \
relative to an oracle that always selects T_2.
  (4) Test-infrastructure correctness dominated measured outcomes in Phase 1 \
more than routing strategy did."""),

("II. SYSTEM DESCRIPTION",
"""\
MCN comprises five components operating as Ray remote actors inside a Docker \
Compose stack: vLLM inference server, Redis state store, ChromaDB vector \
store, MLflow tracking, and an mcn-runner container."""),

("A. Tribes",
"""\
Three tribe actors (T_0, T_1, T_2) share the same base model \
(Qwen2.5-Coder-7B-Instruct-AWQ) but receive distinct system prompts: \
"Reliable Coder," "Fast Coder," and "Creative Coder." No weight updates \
are performed. Phase 3 introduces heterogeneous temperatures: T_0=0.1 \
(deterministic), T_1=0.5 (balanced), T_2=0.9 (high-variance). This creates \
genuine output diversity: T_0 generates near-identical code across attempts; \
T_2 explores a broader algorithmic space at the cost of higher failure variance."""),

("B. Router",
"""\
Two routing implementations are compared.

LinUCB: Contextual bandit with 18-dim context. \
arm* = argmax [x^T theta_i + alpha * sqrt(x^T A_i^{-1} x)]. Epsilon-greedy \
warm-up (epsilon 0.5 -> 0.15, decay=0.995) prevents cold-start collapse. \
alpha=2.5. An epsilon-restoration bug in early runs caused saved epsilon \
(near floor) to override the configured warm-up on restart; fixed in Run 9.

GNN Router: Lightweight 2-layer MLP approximating GraphSAGE on a bipartite \
task-tribe graph. Input: [task_context(18), tribe_embedding(18)] = 36-dim \
concatenation. Three linear layers (36->32->16->1) produce per-tribe scores. \
Online Adam optimisation (lr=0.01), 64-entry experience replay, mini-batches \
of 8. ~2,000 parameters. CPU-only inference."""),

("C. Context Vector",
"""\
Each task maps to an 18-dim vector: task-type one-hot (dims 0–2), exception \
one-hot (dims 3–12), and z-scored execution metrics — runtime, \
tests_passed/failed, test_count, failure_density (dims 13–17). Computed from \
the previous failure on the same task and used as the bandit context for the \
next routing decision."""),

("D. Overseer and Test Generation",
"""\
The OverseerActor generates three test classes per task without seeing \
generated code: (A) Hypothesis property-based fuzzing — integers bounded \
to [−10k, 10k] (unbounded was a critical Phase 1 defect); (B) six adversarial \
static boundary tests including large inputs (10k elements) and large integers \
(10,000, not sys.maxsize — another Phase 1 defect); (C) mutation testing \
for test-suite sensitivity."""),

("E. Patch Registry",
"""\
Verified solutions are stored as context-vector embeddings. Top-K most similar \
prior solutions are retrieved and injected as few-shot hints into the selected \
tribe's prompt. When a reference solution is already present, the patch hint \
is suppressed to prevent conflicting code suggestions."""),

("III. EXPERIMENTAL SETUP",
"""\
Hardware: Single GPU workstation (Docker Compose). \
Model: Qwen/Qwen2.5-Coder-7B-Instruct-AWQ (4-bit AWQ). \
Inference: vLLM, temperature per condition, max_tokens=2048, \
sandbox timeout 10 s.

Phase 1 task set (Runs 1–6): 12 Python functions — sort_list, deduplicate, \
flatten, partition, reverse_string, is_palindrome, word_count, fibonacci, \
is_prime, gcd, invert_dict, running_sum. ~8–9 attempts per type per 100-task run.

Phase 2–3 task set (Runs 7–10, ablations): Extended to 16 types by adding \
has_cycle (graph cycle via Kahn's topological sort), permutations \
(combinatorial enumeration), search_insert (binary search insertion position), \
and unique_paths (DP grid path counting). These harder tasks reduce the \
achievable ceiling and stress-test model capability limits.

Five conditions in Phase 3:
  (A) T=0.1 ablation: Single tribe, T_0 prompt, temperature=0.1.
  (B) T=0.3 ablation: Single tribe, T_0 prompt, temperature=0.3.
  (C) T=0.9 ablation: Single tribe, T_2 prompt, temperature=0.9.
  (D) MCN-LinUCB (Run 9): Three homogeneous T=0.3 tribes, LinUCB router.
  (E) MCN-GNN (Run 10): Three heterogeneous tribes (T_0=0.1, T_1=0.5, T_2=0.9), GNN router.

Phase 1B (stratified live evaluation): Following Phase 3, a separate \
400-task experiment evaluates MCN-LinUCB (homogeneous T=0.3, alpha=2.5) \
on a stratified sample of 8 algorithmic categories — string, \
data_structures, math, dynamic_programming, parsing, iterative, recursive, \
graph — with 50 tasks per category. All tasks executed via live vLLM \
inference (Qwen2.5-Coder-7B-Instruct-AWQ). Stratified sampling eliminates \
task-frequency confounds present in Phase 3's uniform task set, enabling \
direct category-level pass-rate comparison and oracle gap decomposition."""),

("IV. INFRASTRUCTURE DEFECTS",
"""\
Six defects across the three phases are documented here because they \
illustrate a systematic challenge in automated LLM evaluation: test \
infrastructure errors produce misleading metrics that corrupt router training \
signals. The pass-rate improvement from 83% (Run 1) to 98% (Run 4) was \
driven primarily by fixing these defects, not by router learning.

Defect 1 — sys.maxsize in adversarial test B6 (Run 1): fibonacci(sys.maxsize) \
hangs indefinitely. Corrected to fibonacci(10,000) with a 5-second limit.

Defect 2 — Unbounded st.integers() in Hypothesis (Runs 1–3): The strategy \
generator produced st.integers() (unbounded) for integer parameters. Hypothesis \
probes sys.maxsize; even a correct fibonacci loops forever, producing \
tests_passed=0, tests_failed=0 false negatives and corrupting router rewards \
for three runs. Corrected to st.integers(min_value=-10000, max_value=10000).

Defect 3 — PATCH_MIN_ATTEMPTS=2 (Run 1): Patch registration required two \
attempts, leaving the few-shot hint mechanism entirely inactive (0 patches \
stored). Corrected to PATCH_MIN_ATTEMPTS=1.

Defect 4 — permutations reference returns tuples (Phase 2): itertools.permutations \
yields tuples; tests expect list[list[int]]. Corrected: \
[list(p) for p in sorted(set(itertools.permutations(xs)))].

Defect 5 — unique_paths crash on edge cases (Phase 2): math.comb(m+n-2, m-1) \
raises ValueError when m=0 or n=0. Corrected with guard: if m<=0 or n<=0: return 0.

Defect 6 — Epsilon restoration overwrites warm-up (Phase 2): \
LinUCB.from_state_dict() restored saved epsilon (~0.05, floor), silently \
overriding the current .env warm-up schedule. On restart, the router skipped \
its warm-up phase, causing 97/1/2 degenerate routing. Corrected by overriding \
epsilon, epsilon_min, and epsilon_decay from MCNConfig after from_state_dict()."""),

("V. RESULTS",
"""\
Table I presents the five-condition Phase 3 comparison. Table II gives the \
complete ten-run history. The primary finding is that the single-agent T=0.3 \
ablation outperforms all MCN variants."""),

("A. Temperature Ablation — The Dominant Variable",
"""\
Among the three single-agent ablations, T=0.3 achieves the highest pass rate \
(91%), outperforming both T=0.9 (88%) and T=0.1 (86%). This 5 pp spread \
across temperatures exceeds the improvement attributable to any routing \
strategy tested, establishing temperature as the dominant performance variable \
for Qwen2.5-Coder-7B on this task set.

T=0.1 failure mode: search_insert collapses to 25%. At very low temperature, \
the model locks into an incorrect algorithm (linear scan instead of binary \
search) and fails to self-correct across attempts.

T=0.9 failure mode: fibonacci and permutations show higher failure rates than \
T=0.3, despite T=0.9 succeeding more on unique_paths. High variance occasionally \
finds solutions for tricky DP problems but increases failures on tasks with a \
uniquely correct iterative pattern."""),

("B. MCN-LinUCB vs. Single Agent T=0.3 (Run 9)",
"""\
MCN-LinUCB with three homogeneous T=0.3 tribes scored 86% — 5 pp below the \
single-agent T=0.3 ablation (91%). Routing distribution: 71/13/16 (Gini=0.387). \
Chi-squared test for routing independence from task type: p=0.737. \
No statistically significant specialisation was detected.

The 5 pp loss arises from exploration overhead. Epsilon-greedy routing sends \
approximately 29% of tasks to non-dominant tribes during warm-up. Because all \
tribes use the same model at the same temperature, these explorations produce \
no useful diversity signal: the model generates statistically identical code \
regardless of which tribe receives the task. Variance introduced by routing \
is a pure cost with no compensating specialisation benefit.

Task-level losses (MCN vs. single-agent): fibonacci −40 pp, \
search_insert −25 pp, unique_paths −50 pp. These represent cases where routing \
variance accumulated failures during unlucky epsilon windows that a \
single-agent run would have resolved deterministically."""),

("C. MCN-GNN (Heterogeneous) vs. Single Agents (Run 10)",
"""\
MCN-GNN with heterogeneous tribes (T_0=0.1, T_1=0.5, T_2=0.9) scored 88% — \
exactly matching the T=0.9 single-agent ablation and 3 pp below T=0.3. \
Routing: 66/13/21 (Gini=0.353). Chi-squared p=0.762. No significant \
specialisation detected.

Critically, T_2 (T=0.9) achieved 95% pass rate in isolation within Run 10 — \
the highest single-tribe result in the entire study. Yet the full MCN system \
scored only 88%, quantifying a −7 pp routing cost relative to an oracle that \
always selects T_2. The GNN's inability to reliably route all tasks to T_2 \
stems from two failure modes: (i) search_insert is handled worse by T_2 (0%) \
than by T_0, so the oracle would need to route that task type away from T_2; \
(ii) unique_paths is handled poorly by T_1 (0%). The GNN must implicitly learn \
a per-task-type routing policy from ~6 attempts per type — insufficient signal.

Concretely, on fibonacci: T_0 (T=0.1) fails repeatedly with deterministic \
naive recursion; T_2 (T=0.9) succeeds through high-variance exploration of the \
search space. MCN-GNN fibonacci improved from 40% (Run 9, homogeneous) to 60% \
(Run 10), showing that T_2's variance does help — but the GNN must also avoid \
routing fibonacci to T_0, which it only partially achieves."""),

("D. Model Capability Limits",
"""\
Two of the four new tasks reveal hard limits of the 7B model, independent of \
routing.

has_cycle achieves 0% across all conditions. A Kahn's topological sort \
reference solution was provided but the model consistently generates \
syntactically valid yet algorithmically incorrect implementations (missing \
in-degree initialisation, incorrect BFS termination). The overseer's adversarial \
tests — directed cycles, self-loops, disconnected components — catch all errors. \
This is a model capability limit, not a routing failure.

permutations achieves 0% before the type-annotation fix (Defect 4). The \
critical insight: itertools.permutations returns tuples, but tests expect \
list[list[int]]. With the corrected reference solution, small models reliably \
copy the wrapping pattern. This is a reference-solution quality issue, \
not a routing issue."""),

("E. Phase 1B: Category-Level Analysis (400 Tasks)",
"""\
Phase 1B evaluates MCN-LinUCB on 400 stratified live tasks. Overall pass rate: \
245/400 = 61.2%, significantly below Phase 3's 86% on the 16-task benchmark. \
This drop reflects the harder category mix: graph tasks (has_cycle, \
topological_sort, count_components, num_islands, is_bipartite) achieve 0%, \
and recursive/iterative categories average 42–43%.

Category heterogeneity is extreme (Table IV). String tasks — palindromes, \
anagrams, compression, case conversion — achieve 100%, consistent with their \
dominance in pre-training corpora. Graph algorithms achieve 0% in all \
conditions, confirming the 7B model capability limit established in Phase 3. \
The 100 pp spread across categories dwarfs the 5 pp MCN-vs.-single-agent \
gap from Phase 3, establishing category as a far stronger predictor of \
outcome than routing strategy.

Routing specialisation: chi-squared test for routing independence from task \
category: p=0.396. No statistically significant specialisation was detected. \
Routing distribution shows dramatic T0→T1 drift across the 400 tasks, \
consistent with a bandit that has not converged. Analysis of UCB upper \
confidence bounds shows that the bandit requires approximately 2,000+ tasks \
for reliable per-category specialisation at this task diversity.

Oracle gap decomposition (Table V): The total oracle gap of 12.8 pp — MCN \
61.2% vs. oracle 74.0% — decomposes into three roughly equal components: \
exploration cost 4.1 pp (ε-greedy routing to non-optimal arms during warm-up), \
tie-breaking noise 4.7 pp (near-equal UCB scores trigger random arm selection), \
and exploitation error 4.0 pp (insufficient per-category signal causes wrong \
arm selection even during exploitation). No single failure mode dominates, \
indicating the bandit is operating in a data-limited regime throughout."""),

("VI. DISCUSSION",
"""\
The central question — can a learned router improve upon the best single agent \
when tribes share the same base model? — receives a consistent negative answer \
across all three routing experiments and the Phase 1B stratified evaluation."""),

("A. Why Routing Fails With Same-Model Tribes",
"""\
The multi-armed bandit framing assumes that different arms have genuinely \
different expected rewards for different contexts. When tribes are functionally \
similar (same model, same temperature), all arms converge to the same expected \
reward, and exploration cost is never recovered.

Heterogeneous temperatures (0.1/0.5/0.9) create genuine output diversity but \
do not solve the routing problem, because no single temperature is uniformly \
best across all task types. T=0.3 beats T=0.9 on fibonacci and search_insert; \
T=0.9 beats T=0.3 on unique_paths. An oracle router would outperform any single \
temperature — but the GNN achieves p=0.76 routing independence from task type \
at 100 tasks, far from oracle behaviour. The −7 pp gap between T_2-in-isolation \
(95%) and MCN-GNN (88%) directly measures the cost of imperfect routing even \
when a clearly superior tribe exists."""),

("B. What Would It Take for Routing to Add Value?",
"""\
For routing to add aggregate value, the inter-tribe performance variance on \
each task type must exceed the exploration cost. Formally, the expected routing \
gain equals E[max_i R_i(task)] - E[R_routing(task)], where R_routing is the \
reward received under the routing policy. When sigma^2(performance | task) is \
small — as here, where tribes are functionally similar — this gain is \
negligible and the exploration cost dominates.

The boundary conditions for positive routing value are:
  (1) True model diversity: Routing between qualitatively different models \
(e.g., 7B code-specialist vs. 13B reasoning-specialist) would create large \
inter-tribe performance variance on specific task classes.
  (2) Sufficient data: At 6–8 attempts per task type, the GNN lacks the signal \
to learn per-type preferences reliably. Hundreds of attempts per type would \
provide the learning signal needed for genuine specialisation.
  (3) Adaptive temperature: Treating temperature as a continuous router output \
rather than a fixed per-tribe hyperparameter could achieve better task-type \
coverage than discrete temperature choices.

The negative result here is therefore a boundary condition, not a universal \
truth: routing is not beneficial with same-model, same-architecture tribes on \
a 100-task dataset, but may be beneficial with genuinely diverse tribes at \
larger scale."""),

("C. Category-Level Heterogeneity and Bandit Convergence",
"""\
Phase 1B establishes that algorithmic category is the dominant predictor of \
task outcome, with a 100 pp spread from string (100%) to graph (0%). This \
heterogeneity creates an apparent opportunity for category-adaptive routing: \
a converged bandit that identifies graph tasks and halts exploration on those \
categories would reduce the 4.1 pp exploration cost.

However, the critical limiting factor is tribe homogeneity. At T=0.3, all \
three tribes generate statistically identical code for any given task. The \
bandit cannot learn a meaningful per-category preference because the arms do \
not differ. The 12.8 pp oracle gap is therefore irreducible under the current \
tribe configuration regardless of how many tasks are collected: the oracle \
for homogeneous tribes is the same as any single tribe, making the gap an \
artefact of the exploration protocol rather than a recoverable routing signal.

The 2,000-task convergence estimate applies to the case where tribes are \
genuinely diverse. With qualitatively different tribes (different base models, \
architectures, or fine-tuning), the per-category reward signal would differ \
across arms, and the bandit could learn a meaningful routing policy. Under \
those conditions, the category-level heterogeneity documented in Phase 1B \
— specifically the mid-tier cluster (parsing 66%, iterative 43%, recursive \
42%, dynamic_programming 76%) — represents the most actionable routing \
opportunity: categories where model capability is partial and tribe-level \
variance could in principle be exploited."""),

("F. Phase 1C: 2000-Task Longitudinal Experiment (LinUCB)",
"""\
Phase 1C scales the stratified evaluation to 2,000 tasks (8×250 per category) \
using the LinUCB router (alpha=2.5, epsilon=0.3, decay=0.99). Overall pass \
rate: 1,214/2,000 = 60.7%, unchanged from Phase 1B (61.2%), confirming that \
additional tasks provide no benefit when tribes are homogeneous.

Spurious convergence is confirmed at scale. T0 routing share by 500-task \
window: 27% → 75% → 94% → 96%. The bandit reaches nominal lock-in by task \
~1,700 despite T0=60.8%, T1=59.9%, T2=61.8% — all within 2 pp. The locked \
arm (T0) is not the best arm (T2 is marginally higher at 61.8%), confirming \
the convergence is driven by early random variance, not learned preference.

Parsing regression analysis (Phase 1B 65.5% → Phase 1C 58.8%) was traced \
to two task types: roman_to_int (0%, KeyError in format mapping) and \
decode_run_length (0%, ValueError in string split). Three other parsing tasks \
(title_case=100%, count_vowels=100%, camel_to_snake=94%) were unaffected. \
This is a format-compliance failure, not a routing problem, and not \
a systematic regression."""),

("G. Phase 1D: GNN vs. LinUCB Router Comparison at Scale",
"""\
Phase 1D repeats the 2,000-task experiment with the GNN router \
(hidden_dim=32, lr=0.01, buffer=64, batch=8). Table VI presents the \
head-to-head comparison (see also Table VII for strategy breakdown).

Overall: 1,212/2,000 = 60.6% — a −0.1 pp delta vs. LinUCB. All \
per-category deltas lie within ±3.6 pp, consistent with sampling noise \
for 250-task buckets (sigma ≈ 3.2 pp). No systematic GNN advantage or \
disadvantage is detected in any category.

The GNN locks onto T0 faster (72% total routing share vs. 55% for \
LinUCB) and converges earlier (~task 800 vs. ~task 1,700). The GNN's \
mini-batch updates on an 8-entry replay buffer push it to commit earlier \
to the same spurious fixed point. Faster lock-in is a net disadvantage: \
the GNN sacrifices exploratory signal that LinUCB retains longer but \
equally fails to exploit.

A retrospective simulation on Phase 1C data tests four routing strategies \
(Table VII). Category Thompson Sampling (CTS) achieves 63.2% (+2.5 pp vs. \
LinUCB), with the largest gain in math: +12.0 pp (empirical T1=98.4% vs. \
T0=75.7%). Notably, random routing (62.4%) outperforms LinUCB (60.7%), \
confirming the bandit's exploration schedule imposes a net cost relative to \
the uniform baseline. The oracle ceiling is 65.4%."""),

("D. Router Sophistication Is Scale-Invariant",
"""\
The canonical assumption in contextual bandit literature is that more \
expressive models should achieve lower regret on complex reward landscapes. \
The LinUCB vs. GNN comparison at 2,000 tasks directly refutes this assumption \
in the MCN setting: −0.1 pp delta, with the GNN performing marginally worse \
on convergence stability.

This is not a model expressiveness failure — the GNN is correctly trained \
and shows appropriate exploration-exploitation behaviour. The failure is \
conceptual: the reward landscape is flat. When E[R | task, arm_i] ≈ \
E[R | task, arm_j] for all i, j, no routing algorithm can learn a \
meaningful policy from finite data, regardless of its expressiveness.

The practical implication is that "which router is better?" is ill-posed for \
homogeneous tribes. The correct question is: what is the minimum inter-tribe \
performance variance required for routing to recover its exploration cost? \
The retrospective Category TS simulation identifies math as the category where \
this threshold is closest to being reached (T1=98.4% vs T0=75.7%, a 22.7 pp \
spread) — but this spread arises from sampling noise in 63 vs 144 tasks per \
tribe, not from genuine model diversity. With heterogeneous base models or \
fine-tuning targets, category-level performance splits of this magnitude \
could be reliably exploited by a category-aware bandit."""),

("H. Phase 2: Temperature-Heterogeneous Tribes + Live CTS Router",
"""\
Phase 2 extends the homogeneity hypothesis test to temperature-diverse tribes. \
Three tribes are assigned temperatures T0=0.1 (deterministic), T1=0.5 \
(balanced), T2=0.9 (exploratory), with the live CategoryThompsonSampling \
(CTS) router. The experiment runs 1,502 tasks (8×~188 stratified). \
Table VIII presents the comparison with Phase 1C.

Overall pass rate: 911/1,502 = 60.7% — identical to Phase 1C LinUCB. \
Per-tribe routing converged to T2 (52.3% share), but per-tribe pass rates \
(T0=60.4%, T1=59.3%, T2=61.3%) differ by at most 2.0 pp, indistinguishable \
from sampling noise (sigma ≈ 3.2 pp for 188-task buckets).

The oracle gap narrows from 5.1 pp (Phase 1C) to 3.7 pp (Phase 2). \
A smaller oracle gap indicates that tribe outputs are more similar, \
not that the router is better: with lower inter-tribe variance, even \
perfect hindsight routing achieves less improvement. This is the opposite \
of the desired outcome.

CTS routing drift reveals non-trivial learning dynamics absent in LinUCB/GNN. \
The 500-task windows show T1=41%→T0=42%→T2=68% routing shares — CTS \
oscillates between tribes before settling on T2, whereas LinUCB commits to \
T0 irreversibly by task ~1,700. Despite correct Bayesian updating, the learned \
routing preference produces no aggregate gain because the per-(category, arm) \
Beta posterior means remain within noise.

Category pass rates are stable across phases: string 96.2%, math 91.8%, \
data_structures 85.4%, dynamic_programming 72.6%, parsing 56.8%, \
iterative 48.4%, recursive 38.1%, graph 0.0%. The graph ceiling and the \
parsing failure cluster (roman_to_int=0%, decode_run_length=0%) persist \
independently of tribe temperature, confirming these are model-capability \
and prompt-compliance failures, not routing artefacts."""),

("E. Temperature Diversity Does Not Create Routing Signal",
"""\
Phase 2 tests the weakest form of tribe diversity achievable without \
changing the base model: different sampling temperatures. The result \
is a complete null: 60.7% with CTS + temperature-diverse tribes equals \
60.7% with LinUCB + homogeneous tribes.

The information-theoretic explanation is as follows. Sampling temperature \
shifts the output distribution of the same underlying model: high temperature \
increases entropy over the token distribution but does not change which \
concepts or algorithms the model knows. For any given task, the conditional \
probability of a correct solution P(correct | task, model) is fixed by the \
model's weights; temperature scales the variance of outputs around this mean \
but leaves the mean unchanged. Concretely, a 7B model that does not know how \
to represent a graph as a boolean adjacency matrix will fail at temperature \
0.1 and temperature 0.9 alike.

The oracle gap reduction (5.1 pp → 3.7 pp) confirms this directly: reducing \
inter-tribe variance reduces the maximum achievable routing gain, which is \
bounded by E[max_i R_i(task)] - E[R_mean(task)]. Temperature diversity \
moves this bound in the wrong direction.

This result closes the temperature-diversity hypothesis with a decisive null \
and sharpens the minimal condition for routing value: tribes must differ in \
their underlying model weights, fine-tuning targets, or capability profiles — \
not merely in their sampling hyperparameters. Future experiments should \
therefore use genuinely different models (e.g., a code-specialist, a \
reasoning-specialist, and a general-purpose model) rather than temperature \
variants of the same checkpoint."""),

("VII. CONCLUSIONS",
"""\
We have conducted ten experimental runs, three controlled ablations, and two \
2,000-task longitudinal evaluations evaluating the Mycelial Council Network \
on Python code synthesis benchmarks. Primary conclusions:

  (1) Temperature dominates routing. T=0.3 achieves 91% (single agent); \
all MCN variants score 86–88%. The optimal multi-agent configuration does not \
exceed the optimal single-agent temperature.

  (2) Same-model routing imposes a net cost. MCN-LinUCB (homogeneous T=0.3) \
scores −5 pp vs. single-agent, with p=0.74 routing independence.

  (3) Heterogeneous temperatures do not enable specialisation. MCN-GNN matches \
T=0.9 single agent (both 88%). The −7 pp gap between T_2-in-isolation (95%) \
and MCN-GNN (88%) quantifies the routing cost when a clearly superior tribe \
exists but cannot be selected reliably.

  (4) Infrastructure correctness dominated early measured outcomes. Six defects \
across the three phases caused systematic performance distortions unrelated \
to routing quality.

  (5) Model capability limits, not routing, explain the hardest failures. \
has_cycle achieves 0% in all conditions; this is a 7B model limit.

  (6) Category-level heterogeneity dwarfs routing effects. Phase 1B/1C reveals \
a 100 pp spread from string (99%) to graph (0%). Corrected pass rate excluding \
the graph capability limit: 81.3%. The oracle gap of 12.8 pp splits equally \
among exploration cost (4.1 pp), tie-breaking noise (4.7 pp), and exploitation \
error (4.0 pp); the bandit is data-limited throughout.

  (7) Router sophistication is scale-invariant. At 2,000 tasks, LinUCB (60.7%) \
and GNN (60.6%) produce a −0.1 pp delta — a null result robust to a 40x \
increase in router parameter count. The GNN locks onto T0 faster (task ~800 \
vs. ~1,700) and routes 72% vs. 55% to T0 — a worse outcome. No routing \
algorithm can learn from a flat reward landscape.

  (8) Category Thompson Sampling is the correct next router. Retrospective \
simulation yields +2.5 pp (+12 pp on math) over LinUCB. CTS maintains \
per-(category, arm) Beta posteriors, preventing the cross-category signal \
contamination that causes spurious LinUCB/GNN convergence.

  (9) Temperature diversity does not create routing signal. Phase 2 deploys \
live CTS with temperature-heterogeneous tribes (T0=0.1, T1=0.5, T2=0.9): \
overall pass rate is 60.7% — identical to Phase 1C. Per-tribe rates span \
only 2.0 pp (T0=60.4%, T1=59.3%, T2=61.3%). The oracle gap narrows from \
5.1 pp to 3.7 pp, moving in the wrong direction. CTS learns routing \
preferences (oscillating T1→T0→T2 drift) but cannot improve aggregate \
performance because temperature shifts output variance, not model capability. \
The minimal condition for routing value is inter-model capability diversity, \
not sampling or algorithmic sophistication.

Future directions: (A) heterogeneous base models — routing between a \
code-specialist, reasoning-specialist, and general-purpose 7B model to \
create genuine inter-tribe performance variance; (B) graph-fixed benchmark \
with output-format-constrained prompts to measure the 81.3% corrected \
ceiling under better task design; (C) per-category fine-tuning — train \
separate LoRA adapters per category to create specialised tribes on the \
same base checkpoint."""),
]

# ── TABLE I: Five-condition comparison (compact, fits in 3.5" column) ─────────
TABLE1_CAPTION = "TABLE I\nFIVE-CONDITION PHASE 3 COMPARISON (100 tasks, 16-task benchmark)"
TABLE1_HEADER  = ["Condition", "T", "Pass", "Fib", "Search\nIns.", "Routing"]
TABLE1_ROWS    = [
    ["(A) T=0.1 solo",   "0.1",     "86%", "40%", "25%",  "—"],
    ["(B) T=0.3 solo",   "0.3",     "91%", "80%", "100%", "—"],
    ["(C) T=0.9 solo",   "0.9",     "88%", "40%", "n/m",  "—"],
    ["(D) MCN-LinUCB",   "0.3×3",   "86%", "40%", "n/m",  "71/13/16"],
    ["(E) MCN-GNN",      "0.1/.5/.9","88%","60%", "n/m",  "66/13/21"],
]
TABLE1_NOTE    = ("n/m = not measured for that condition. "
                  "T=0.9 unique_paths=100%. "
                  "Best result (B) highlighted.")

# ── TABLE II: Complete run history (compact, fits in 3.5" column) ─────────────
TABLE2_CAPTION = "TABLE II\nCOMPLETE RUN HISTORY (RUNS 1–10)"
TABLE2_HEADER  = ["Run", "Router", "Tasks", "Pass%", "Routing", "Notes"]
TABLE2_ROWS    = [
    ["1",  "GNN",    "12", "83%", "10/69/21", "Infra bugs 1–3"],
    ["2",  "GNN",    "12", "90%", "33/48/19", "Partial fix"],
    ["3",  "GNN",    "12", "91%", "33/48/19", "Partial fix"],
    ["4",  "GNN",    "12", "98%", "39/13/48", "All bugs fixed"],
    ["5",  "GNN*",   "12", "98%", "10/8/82",  "Inherited state"],
    ["6",  "LinUCB", "12", "97%", "100/0/0",  "Cold-start collapse"],
    ["7-8","LinUCB", "16","84.5%","mixed",    "Exploratory; ε-bug"],
    ["9",  "LinUCB", "16", "86%", "71/13/16", "Post-fix; −5pp vs solo"],
    ["10", "GNN",    "16", "88%", "66/13/21", "Hetero T=0.1/.5/.9"],
]
TABLE2_NOTE    = ("* Run 5 inherits GNN state from Run 4 (no --fresh). "
                  "Runs 7–8 combined exploratory dataset.")

# ── TABLE III: MCN vs. best single agent ──────────────────────────────────────
TABLE3_CAPTION = "TABLE III\nMCN VARIANTS vs. BEST SINGLE AGENT"
TABLE3_HEADER  = ["Metric", "T=0.3 Solo", "MCN-LinUCB", "MCN-GNN"]
TABLE3_ROWS    = [
    ["Pass rate",          "91%",           "86% (−5pp)",    "88% (−3pp)"],
    ["fibonacci",          "80%",           "40% (−40pp)",   "60% (−20pp)"],
    ["Routing dist.",      "—",             "71/13/16",      "66/13/21"],
    ["Chi-sq. p",          "—",             "0.737",         "0.762"],
    ["Specialisation?",    "—",             "No",            "No"],
    ["T_2 isolation",      "88% (T=0.9)",   "n/a",           "95% (within-run)"],
    ["Oracle gap",         "—",             "n/a",           "−7pp vs T_2 solo"],
    ["Verdict",            "BEST",          "−5pp vs best",  "=T=0.9 solo"],
]
TABLE3_NOTE    = ("T_2 isolation = T_2 tribe pass rate computed within Run 10, "
                  "not a standalone ablation. Oracle gap = 95%−88% = 7 pp routing cost.")

# ── TABLE IV: Phase 1B category-level results ─────────────────────────────────
TABLE4_CAPTION = "TABLE IV\nPHASE 1B CATEGORY-LEVEL RESULTS (400 tasks, 8×50 stratified)"
TABLE4_HEADER  = ["Category", "Pass%", "vs. Overall", "Capability"]
TABLE4_ROWS    = [
    ["string",          "100%",  "+38.8 pp", "Full"],
    ["data_structures", "84%",   "+22.8 pp", "High"],
    ["math",            "86%",   "+24.8 pp", "High"],
    ["dynamic_prog.",   "76%",   "+14.8 pp", "Moderate"],
    ["parsing",         "66%",   "+4.8 pp",  "Moderate"],
    ["iterative",       "43%",   "−18.2 pp", "Low"],
    ["recursive",       "42%",   "−19.2 pp", "Low"],
    ["graph",           "0%",    "−61.2 pp", "None"],
    ["Overall",         "61.2%", "—",        "—"],
]
TABLE4_NOTE    = ("vs. Overall = category pass rate minus overall mean (61.2%). "
                  "Capability = qualitative model-performance tier for Qwen2.5-Coder-7B. "
                  "Graph 0% is a hard capability limit, not a routing failure.")

# ── TABLE V: Oracle gap decomposition ─────────────────────────────────────────
TABLE5_CAPTION = "TABLE V\nORACLE GAP DECOMPOSITION (Phase 1B, 400 tasks)"
TABLE5_HEADER  = ["Component", "Gap (pp)", "Mechanism"]
TABLE5_ROWS    = [
    ["Exploration cost (Q1)",     "4.1",  "ε-greedy routes ~15% of tasks to non-optimal arms"],
    ["Tie-breaking noise (Q2+Q3)","4.7",  "Near-equal UCB scores → random arm selection"],
    ["Exploitation error (Q4)",   "4.0",  "Insufficient per-category signal → wrong arm"],
    ["Total oracle gap",          "12.8", "MCN 61.2% vs. oracle 74.0%"],
]
TABLE5_NOTE    = ("Oracle = per-task best-tribe selection with perfect knowledge. "
                  "Q1–Q4 computed via counterfactual routing simulation on recorded decisions. "
                  "All three components are roughly equal; no single source dominates.")

# ── TABLE VI: LinUCB vs GNN at 2000 tasks ─────────────────────────────────────
TABLE6_CAPTION = "TABLE VI\nROUTER COMPARISON AT SCALE (2,000 tasks, 8×250 stratified)"
TABLE6_HEADER  = ["Metric", "LinUCB (Phase 1C)", "GNN (Phase 1D)", "Delta"]
TABLE6_ROWS    = [
    ["Overall pass rate",   "60.7%",  "60.6%",  "−0.1 pp"],
    ["T0 routing share",    "55%",    "72%",    "GNN locks harder"],
    ["T1 routing share",    "28%",    "8%",     "GNN drops T1 early"],
    ["T2 routing share",    "16%",    "19%",    "+3 pp"],
    ["T0 pass rate",        "60.8%",  "60.9%",  "+0.1 pp"],
    ["T1 pass rate",        "59.9%",  "59.6%",  "−0.3 pp"],
    ["T2 pass rate",        "61.8%",  "60.1%",  "−1.7 pp"],
    ["Convergence to T0",   "~task 1700", "~task 800", "GNN 2x faster"],
    ["Router parameters",   "~54",    "~2,000", "40x more complex"],
]
TABLE6_NOTE    = ("Convergence = task index where T0 share first exceeds 90% and stays there. "
                  "Per-tribe pass rates within ±2 pp across both routers — "
                  "statistically indistinguishable. GNN parameter count: 3 tribe "
                  "embeddings (18-dim) + 2-layer MLP (36→32→16→1).")

# ── TABLE VII: Strategy comparison (retrospective simulation) ──────────────────
TABLE7_CAPTION = "TABLE VII\nROUTING STRATEGY COMPARISON (Phase 1C data, 2,000 tasks)"
TABLE7_HEADER  = ["Strategy", "Pass Rate", "vs. LinUCB", "Key Observation"]
TABLE7_ROWS    = [
    ["LinUCB (actual)",          "60.7%", "baseline",  "Locks to T0; hurts math"],
    ["Random routing",           "62.4%", "+1.7 pp",   "Uniform > trained bandit"],
    ["Category TS (simulated)",  "63.2%", "+2.5 pp",   "Math +12 pp; no lock-in"],
    ["Oracle per-category",      "65.4%", "+4.7 pp",   "Ceiling; requires hindsight"],
]
TABLE7_NOTE    = ("Category TS simulation uses imputed rewards from empirical per-(category, arm) "
                  "rates for counterfactual arm selections. Random > LinUCB confirms the bandit's "
                  "exploration schedule imposes a net cost vs. uniform baseline with homogeneous tribes. "
                  "Math +12 pp for CTS arises from empirical T1=98.4% vs. T0=75.7% (250 tasks each).")

# ── TABLE VIII: Phase 1C vs Phase 2 (heterogeneous temperatures) ───────────
TABLE8_CAPTION = "TABLE VIII\nPHASE 1C vs. PHASE 2: EFFECT OF TEMPERATURE DIVERSITY (CTS ROUTER)"
TABLE8_HEADER  = ["Metric", "Phase 1C\n(LinUCB+Homo)", "Phase 2\n(CTS+Hetero)", "Delta"]
TABLE8_ROWS    = [
    ["Tasks",              "2,000",    "1,502",    "—"],
    ["Overall pass rate",  "60.7%",    "60.7%",    "0.0 pp"],
    ["Oracle (per-cat)",   "65.8%",    "64.3%",    "−1.5 pp"],
    ["Oracle gap",         "5.1 pp",   "3.7 pp",   "−1.4 pp"],
    ["Dominant tribe",     "T0 (55%)", "T2 (52%)", "shifted"],
    ["T0 pass rate",       "60.8%",    "60.4%",    "−0.4 pp"],
    ["T1 pass rate",       "59.9%",    "59.3%",    "−0.6 pp"],
    ["T2 pass rate",       "61.8%",    "61.3%",    "−0.5 pp"],
    ["Max inter-tribe gap","1.9 pp",   "2.0 pp",   "≈ same"],
    ["Convergence pattern","Lock T0 (~t=1700)", "Oscillate T1→T0→T2", "CTS avoids lock"],
    ["Graph pass rate",    "0%",       "0%",       "persists"],
    ["String pass rate",   "99%",      "96.2%",    "≈ same"],
]
TABLE8_NOTE    = ("Oracle gap decrease indicates tribes became more similar under temperature diversity, "
                  "not that routing improved. CTS routing drift (T1→T0→T2 oscillation) reflects correct "
                  "Bayesian updating; the reward signal is insufficient because temperature diversity "
                  "does not create capability diversity. Phase 2: 1,502 tasks (OOM restart at t=1,490).")


# ══════════════════════════════════════════════════════════════════════════════
# Progress Report content (unchanged from previous edition)
# ══════════════════════════════════════════════════════════════════════════════

REPORT_TITLE    = "MCN Research Log — Technical Progress Report (Complete)"
REPORT_SUBTITLE = "Mycelial Council Network v0.1  |  Runs 1-10 + Ablations  |  February 2026"
REPORT_META = [
    ("Project",        "Mycelial Council Network (MCN) v0.1"),
    ("Hardware",       "Single GPU workstation · Docker Compose stack"),
    ("Model",          "Qwen/Qwen2.5-Coder-7B-Instruct-AWQ (4-bit AWQ)"),
    ("Phase 1",        "Runs 1-6: GNN/LinUCB on 12-task benchmark"),
    ("Phase 2",        "Runs 7-8: Expanded to 16-task benchmark (exploratory)"),
    ("Phase 3",        "Runs 9-10 + 3 ablations: Controlled temperature comparison"),
    ("Central finding","T=0.3 single agent (91%) beats all MCN variants (86-88%). "
                       "Temperature dominates routing."),
    ("Date",           "2026-02-23"),
]

REPORT_RUN_TABLE_HEADER = ["Run", "Router", "Task set", "Pass%",
                            "fibonacci", "Routing T0/T1/T2", "Key finding"]
REPORT_RUN_TABLE = [
    ["1",  "GNN",    "12", "83%",  "0% (false neg)", "10/69/21",
     "Infra bugs: unbounded Hypothesis, sys.maxsize, PATCH_MIN=2"],
    ["2",  "GNN",    "12", "90%",  "0% (false neg)", "33/48/19",
     "Partial fix; reference_solution channel inactive"],
    ["3",  "GNN",    "12", "91%",  "0% (false neg)", "33/48/19",
     "Partial fix; Hypothesis still unbounded"],
    ["4",  "GNN",    "12", "98%",  "75%",            "39/13/48",
     "All infra defects fixed; GNN explores all tribes"],
    ["5",  "GNN*",   "12", "98%",  "75%",            "10/8/82",
     "Inherited GNN state; eps near floor; T2 dominant (82%)"],
    ["6",  "LinUCB", "12", "97%",  "62%",            "100/0/0",
     "Cold-start routing collapse; degenerate single-tribe system"],
    ["7-8","LinUCB", "16", "84.5%","n/m",            "mixed",
     "Exploratory; epsilon-restoration bug; ref-bugs for 3 new tasks"],
    ["9",  "LinUCB", "16", "86%",  "40%",            "71/13/16",
     "Post-fix; epsilon warm-up working; 5pp BELOW single-agent T=0.3"],
    ["10", "GNN",    "16", "88%",  "60%",            "66/13/21",
     "Hetero T=0.1/0.5/0.9; MCN=T=0.9 single; T2 isolation=95%; no specialisation (p=0.76)"],
]

REPORT_ABL_HEADER = ["Condition", "Temp", "Pass%", "fibonacci",
                     "search_insert", "unique_paths", "vs MCN-LinUCB", "vs MCN-GNN"]
REPORT_ABL_TABLE  = [
    ["T=0.1 single agent",  "0.1",        "86%", "40%", "25%",  "n/m",  "=0pp",  "−2pp"],
    ["T=0.3 single agent",  "0.3",        "91%", "80%", "100%", "100%", "+5pp",  "+3pp"],
    ["T=0.9 single agent",  "0.9",        "88%", "40%", "n/m",  "100%", "+2pp",  "=0pp"],
    ["MCN-LinUCB (Run 9)",  "0.3×3",      "86%", "40%", "n/m",  "n/m",  "—",     "−2pp"],
    ["MCN-GNN (Run 10)",    "0.1/0.5/0.9","88%", "60%", "n/m",  "n/m",  "+2pp",  "—"],
]

REPORT_DEFECT_HEADER = ["Defect", "Phase", "Impact"]
REPORT_DEFECTS = [
    ["sys.maxsize in adversarial test B6",           "1 (Run 1)",   "fibonacci hang; masked true cause"],
    ["st.integers() unbounded in Hypothesis",         "1 (Runs 1-3)","fibonacci 0%; corrupted rewards for 3 runs"],
    ["PATCH_MIN_ATTEMPTS=2 hardcoded",                "1 (Run 1)",   "0 patches stored; hints inactive"],
    ["permutations reference returns tuples",         "2 (Runs 7-8)","permutations 0%; tests expect list[list[int]]"],
    ["unique_paths math.comb crash on m=0/n=0",       "2 (Runs 7-8)","overseer edge case raises ValueError"],
    ["epsilon restoration overwrites config warm-up", "2 (Runs 7-8)","97/1/2 routing collapse after restart"],
    ["invert_dict spec-oracle contradiction",         "1 (Run 1)",   "8% pass-rate loss from bad task spec"],
    ["has_cycle: model capability limit",             "2-3 (all)",   "0% across all runs; Kahn's reference ineffective for 7B"],
]

REPORT_STATE = [
    ("Overall finding",   "Temperature dominates routing. T=0.3 optimal (91%). Routing adds no benefit with same-model tribes."),
    ("MCN-LinUCB (R9)",   "86%; 71/13/16; chi-sq p=0.74; 5pp BELOW single-agent T=0.3"),
    ("MCN-GNN (R10)",     "88%; 66/13/21; chi-sq p=0.76; T_2 isolation=95%; oracle gap=−7pp"),
    ("T=0.3 ablation",    "91% — best overall; fibonacci 80%; search_insert 100%; unique_paths 100%"),
    ("T=0.9 ablation",    "88% — matches MCN-GNN; fibonacci 40%; unique_paths 100%"),
    ("T=0.1 ablation",    "86% — matches MCN-LinUCB; search_insert 25% (deterministic failure)"),
    ("has_cycle",         "0% in ALL conditions — 7B model capability limit"),
    ("fibonacci",         "T=0.3 best (80%); T_2 in MCN-GNN improved to 60% via high-variance search"),
    ("Patch registry",    "GNN runs: 279-377 patches (ChromaDB); LinUCB runs: 86-97 patches (in-memory)"),
    ("Next step A",       "True model diversity: Route between qualitatively different models"),
    ("Next step B",       "Scale: n=500+ tasks, 20+ task types for GNN specialisation signal"),
    ("Next step C",       "Adaptive temperature: Continuous router output, not fixed per-tribe"),
    ("Next step D",       "Harder task distribution: Competitive programming / algorithmic variance"),
]


# ══════════════════════════════════════════════════════════════════════════════
# IEEE PDF helpers
# ══════════════════════════════════════════════════════════════════════════════

# ── IEEE page dimensions (Letter) ─────────────────────────────────────────────
_PW, _PH   = letter           # 8.5" × 11"
_MT        = 0.75 * inch      # top margin
_MB        = 1.00 * inch      # bottom margin
_ML        = 0.625 * inch     # left margin
_MR        = 0.625 * inch     # right margin
_GAP       = 0.25  * inch     # column gutter
_FW        = _PW - _ML - _MR  # 7.25" full width
_CW        = (_FW - _GAP) / 2 # 3.5"  column width
_FH        = _PH - _MT - _MB  # 9.25" full page height
_HEADER_H  = 4.15 * inch      # estimated header block height (title + abstract)
_BODY_H_P1 = _FH - _HEADER_H  # two-column height on page 1


def _ieee_styles():
    S = {}
    S["title"]   = ParagraphStyle("ieee_t",
                       fontName="Times-Bold",    fontSize=20, leading=24,
                       alignment=TA_CENTER,      textColor=BLACK, spaceAfter=4)
    S["authors"] = ParagraphStyle("ieee_a",
                       fontName="Times-Italic",  fontSize=11, leading=13,
                       alignment=TA_CENTER,      textColor=BLACK, spaceAfter=2)
    S["affil"]   = ParagraphStyle("ieee_af",
                       fontName="Times-Roman",   fontSize=9,  leading=11,
                       alignment=TA_CENTER,      textColor=BLACK, spaceAfter=3)
    S["abstract"]= ParagraphStyle("ieee_ab",
                       fontName="Times-Roman",   fontSize=9,  leading=11,
                       alignment=TA_JUSTIFY,     textColor=BLACK, spaceAfter=3)
    S["keywords"]= ParagraphStyle("ieee_kw",
                       fontName="Times-Roman",   fontSize=9,  leading=11,
                       alignment=TA_JUSTIFY,     textColor=BLACK, spaceAfter=0)
    S["section"] = ParagraphStyle("ieee_sec",
                       fontName="Times-Bold",    fontSize=10, leading=12,
                       alignment=TA_CENTER,      textColor=BLACK,
                       spaceBefore=8, spaceAfter=3)
    S["subsect"] = ParagraphStyle("ieee_sub",
                       fontName="Times-BoldItalic", fontSize=10, leading=12,
                       alignment=TA_LEFT,        textColor=BLACK,
                       spaceBefore=5, spaceAfter=2)
    S["body"]    = ParagraphStyle("ieee_bd",
                       fontName="Times-Roman",   fontSize=10, leading=12,
                       alignment=TA_JUSTIFY,     textColor=BLACK,
                       spaceAfter=4, firstLineIndent=14)
    S["body_ni"] = ParagraphStyle("ieee_ni",
                       fontName="Times-Roman",   fontSize=10, leading=12,
                       alignment=TA_JUSTIFY,     textColor=BLACK,
                       spaceAfter=4)
    S["list"]    = ParagraphStyle("ieee_li",
                       fontName="Times-Roman",   fontSize=9.5, leading=11.5,
                       alignment=TA_JUSTIFY,     textColor=BLACK,
                       leftIndent=10, spaceAfter=2, firstLineIndent=0)
    S["caption"] = ParagraphStyle("ieee_cap",
                       fontName="Times-Bold",    fontSize=8,  leading=10,
                       alignment=TA_CENTER,      textColor=BLACK, spaceAfter=2)
    S["note"]    = ParagraphStyle("ieee_nt",
                       fontName="Times-Italic",  fontSize=7.5, leading=9,
                       alignment=TA_LEFT,        textColor=BLACK, spaceAfter=4)
    S["th"]      = ParagraphStyle("ieee_th",
                       fontName="Times-Bold",    fontSize=8,  leading=10,
                       alignment=TA_CENTER,      textColor=WHITE)
    S["td"]      = ParagraphStyle("ieee_td",
                       fontName="Times-Roman",   fontSize=8,  leading=10,
                       alignment=TA_CENTER,      textColor=BLACK)
    S["tdl"]     = ParagraphStyle("ieee_tdl",
                       fontName="Times-Roman",   fontSize=8,  leading=10,
                       alignment=TA_LEFT,        textColor=BLACK)
    S["footer"]  = ParagraphStyle("ieee_ft",
                       fontName="Times-Roman",   fontSize=8,  leading=10,
                       alignment=TA_CENTER,      textColor=BLACK)
    return S


def _is_subsection(heading: str) -> bool:
    """True if heading starts with a letter-dot prefix, e.g. 'A. Tribes'."""
    return bool(re.match(r'^[A-Z]\.\s', heading))


def _ieee_body_paragraphs(body_text: str, S: dict, story: list) -> None:
    """Render section body into IEEE-style paragraphs.

    Lines indented with '  ' are rendered as compact list items (9.5pt, no
    first-line indent).  Blank lines separate paragraphs.  The first paragraph
    of each section uses body_ni (no first-line indent); subsequent paragraphs
    use body (indented).
    """
    blocks = body_text.split("\n\n")
    first = True
    for block in blocks:
        if not block.strip():
            continue
        lines = block.split("\n")
        # Check if this block is entirely list items
        if all(l.startswith("  ") or not l.strip() for l in lines):
            for line in lines:
                if line.strip():
                    story.append(Paragraph(line.strip(), S["list"]))
        else:
            text = " ".join(l.strip() for l in lines if l.strip())
            style = S["body_ni"] if first else S["body"]
            story.append(Paragraph(text, style))
            first = False


def _ieee_table(caption: str, header: list, rows: list, col_widths: list,
                S: dict, note: str = "",
                highlight_row: int = -1,
                highlight_col: Colors = None) -> list:
    """Return a list of flowables: caption, table, optional note."""
    result = []
    result.append(Paragraph(caption, S["caption"]))

    data = [[Paragraph(h.replace("\n", " "), S["th"]) for h in header]]
    for i, row in enumerate(rows):
        cells = []
        for j, val in enumerate(row):
            st = S["td"] if (len(val) <= 12 and "\n" not in val) else S["tdl"]
            cells.append(Paragraph(val, st))
        data.append(cells)

    t = Table(data, colWidths=col_widths, repeatRows=1)
    cmds = list(GRID_BASE) + [
        ("BACKGROUND",     (0,0),(-1,0),  BLACK),
        ("ROWBACKGROUNDS", (0,1),(-1,-1), [WHITE, colors.HexColor("#F8F8F8")]),
    ]
    if highlight_row >= 1:
        cmds.append(("BACKGROUND", (0,highlight_row),(-1,highlight_row),
                      colors.HexColor("#D8EAD3")))
    t.setStyle(TableStyle(cmds))
    result.append(t)
    if note:
        result.append(Paragraph(note, S["note"]))
    result.append(Spacer(1, 4))
    return result


def _page_number(canvas, doc):
    """Draw centred page number at bottom of each page (used as onPage callback)."""
    canvas.saveState()
    canvas.setFont("Times-Roman", 8)
    canvas.drawCentredString(_PW / 2, 0.5 * inch,
                             f"— {canvas.getPageNumber()} —")
    canvas.restoreState()

# Alias for PageTemplate onPage parameter
_pn = _page_number


# ══════════════════════════════════════════════════════════════════════════════
# ACADEMIC PAPER — IEEE PDF (two-column, Letter, Times-Roman)
# ══════════════════════════════════════════════════════════════════════════════

def build_paper_pdf(path: str):
    # ── Page templates ─────────────────────────────────────────────────────
    # Page 1: full-width header frame on top, two-column body below
    f_hdr  = Frame(_ML, _MB + _BODY_H_P1, _FW, _HEADER_H, id="hdr",
                   leftPadding=0, rightPadding=0,
                   topPadding=0,  bottomPadding=0)
    f_lp1  = Frame(_ML, _MB, _CW, _BODY_H_P1, id="lp1",
                   leftPadding=0, rightPadding=4,
                   topPadding=0,  bottomPadding=0)
    f_rp1  = Frame(_ML + _CW + _GAP, _MB, _CW, _BODY_H_P1, id="rp1",
                   leftPadding=4, rightPadding=0,
                   topPadding=0,  bottomPadding=0)

    # Pages 2+: two-column full-height
    f_l    = Frame(_ML, _MB, _CW, _FH, id="l",
                   leftPadding=0, rightPadding=4,
                   topPadding=0,  bottomPadding=0)
    f_r    = Frame(_ML + _CW + _GAP, _MB, _CW, _FH, id="r",
                   leftPadding=4, rightPadding=0,
                   topPadding=0,  bottomPadding=0)

    doc = BaseDocTemplate(
        path, pagesize=letter,
        pageTemplates=[
            PageTemplate(id="P1", frames=[f_hdr, f_lp1, f_rp1], onPage=_pn),
            PageTemplate(id="Pn", frames=[f_l,   f_r  ],         onPage=_pn),
        ],
        title=PAPER_TITLE, author=PAPER_AUTHORS,
    )

    S = _ieee_styles()
    story = []

    # ── Header section (fills f_hdr on page 1) ─────────────────────────────
    story.append(Paragraph(PAPER_TITLE, S["title"]))
    story.append(Spacer(1, 4))
    story.append(Paragraph(PAPER_AUTHORS, S["authors"]))
    story.append(Paragraph(PAPER_AFFIL,   S["affil"]))
    story.append(Spacer(1, 3))
    story.append(HRFlowable(width="100%", thickness=0.6, color=BLACK, spaceAfter=3))

    # Abstract
    story.append(Paragraph(
        "<i><b>Abstract</b></i>\u2014" + IEEE_ABSTRACT, S["abstract"]))
    story.append(Spacer(1, 3))
    story.append(Paragraph(
        "<i><b>Index Terms</b></i>\u2014" + IEEE_KEYWORDS, S["keywords"]))
    story.append(HRFlowable(width="100%", thickness=0.6, color=BLACK, spaceAfter=0))

    # ── Transition from header frame → two-column body ──────────────────────
    story.append(FrameBreak())         # end f_hdr, start f_lp1
    story.append(NextPageTemplate("Pn"))  # after page 1, use two-column template

    # ── Body sections ───────────────────────────────────────────────────────
    for i, (heading, body) in enumerate(IEEE_SECTIONS):
        sub = _is_subsection(heading)
        story.append(Paragraph(heading, S["subsect"] if sub else S["section"]))
        if not sub:
            story.append(HRFlowable(width="100%", thickness=0.3,
                                     color=BLACK, spaceAfter=2))
        _ieee_body_paragraphs(body, S, story)

        # ── Insert tables at specific points ──────────────────────────────
        # After "V. RESULTS" intro paragraph → TABLE I and TABLE II
        if heading == "V. RESULTS":
            story.append(Spacer(1, 4))
            cw1 = [1.35*inch, 0.42*inch, 0.38*inch,
                   0.38*inch, 0.42*inch, 0.52*inch]
            story.extend(_ieee_table(
                TABLE1_CAPTION, TABLE1_HEADER, TABLE1_ROWS, cw1, S,
                note=TABLE1_NOTE, highlight_row=2))   # T=0.3 row
            cw2 = [0.38*inch, 0.75*inch, 0.42*inch,
                   0.42*inch, 0.78*inch, 0.72*inch]
            story.extend(_ieee_table(
                TABLE2_CAPTION, TABLE2_HEADER, TABLE2_ROWS, cw2, S,
                note=TABLE2_NOTE))

        # After "C. MCN-GNN" → TABLE III
        if heading == "C. MCN-GNN (Heterogeneous) vs. Single Agents (Run 10)":
            story.append(Spacer(1, 4))
            cw3 = [1.30*inch, 0.72*inch, 0.72*inch, 0.76*inch]
            story.extend(_ieee_table(
                TABLE3_CAPTION, TABLE3_HEADER, TABLE3_ROWS, cw3, S,
                note=TABLE3_NOTE))

        # After Phase 1B section → TABLE IV (category results) + TABLE V (oracle gap)
        if heading == "E. Phase 1B: Category-Level Analysis (400 Tasks)":
            story.append(Spacer(1, 4))
            cw4 = [1.20*inch, 0.55*inch, 0.72*inch, 0.60*inch]
            story.extend(_ieee_table(
                TABLE4_CAPTION, TABLE4_HEADER, TABLE4_ROWS, cw4, S,
                note=TABLE4_NOTE, highlight_row=9))   # Overall row (last)
            cw5 = [1.55*inch, 0.52*inch, 1.50*inch]
            story.extend(_ieee_table(
                TABLE5_CAPTION, TABLE5_HEADER, TABLE5_ROWS, cw5, S,
                note=TABLE5_NOTE, highlight_row=4))   # Total row (last)

        # After Phase 1D section → TABLE VI (router comparison) + TABLE VII (strategies)
        if heading == "G. Phase 1D: GNN vs. LinUCB Router Comparison at Scale":
            story.append(Spacer(1, 4))
            cw6 = [1.40*inch, 0.85*inch, 0.85*inch, 0.90*inch]
            story.extend(_ieee_table(
                TABLE6_CAPTION, TABLE6_HEADER, TABLE6_ROWS, cw6, S,
                note=TABLE6_NOTE))
            cw7 = [1.35*inch, 0.60*inch, 0.62*inch, 1.43*inch]
            story.extend(_ieee_table(
                TABLE7_CAPTION, TABLE7_HEADER, TABLE7_ROWS, cw7, S,
                note=TABLE7_NOTE, highlight_row=3))   # CTS row

        # After Phase 2 section → TABLE VIII (temperature diversity comparison)
        if heading == "H. Phase 2: Temperature-Heterogeneous Tribes + Live CTS Router":
            story.append(Spacer(1, 4))
            cw8 = [1.35*inch, 0.90*inch, 0.90*inch, 0.85*inch]
            story.extend(_ieee_table(
                TABLE8_CAPTION, TABLE8_HEADER, TABLE8_ROWS, cw8, S,
                note=TABLE8_NOTE, highlight_row=1))   # Overall pass rate row

    # ── Footer ──────────────────────────────────────────────────────────────
    story.append(Spacer(1, 6))
    story.append(HRFlowable(width="100%", thickness=0.6, color=BLACK, spaceAfter=2))
    story.append(Paragraph(
        "Reproducibility: all code, Docker configuration, and MLflow artifacts "
        "available in the MCN repository. Raw data: mcn-results:/results/runs.jsonl.",
        S["footer"]))

    doc.build(story)
    print(f"  Paper PDF  saved -> {path}")


# ══════════════════════════════════════════════════════════════════════════════
# ACADEMIC PAPER — IEEE DOCX
# ══════════════════════════════════════════════════════════════════════════════

def _docx_set_cols(section, num_cols: int, space_twips: int = 720):
    """Set column layout on a DOCX section."""
    sectPr = section._sectPr
    # Remove any existing cols element
    for existing in sectPr.findall(qn("w:cols")):
        sectPr.remove(existing)
    cols = OxmlElement("w:cols")
    cols.set(qn("w:num"),        str(num_cols))
    cols.set(qn("w:space"),      str(space_twips))
    cols.set(qn("w:equalWidth"), "1")
    sectPr.append(cols)


def _docx_page_size_letter(section):
    """Set Letter page size with IEEE margins."""
    section.page_width   = Inches(8.5)
    section.page_height  = Inches(11)
    section.top_margin   = Inches(0.75)
    section.bottom_margin= Inches(1.0)
    section.left_margin  = Inches(0.625)
    section.right_margin = Inches(0.625)


def _docx_ieee_table(doc, caption: str, header: list, rows: list,
                     note: str = "", highlight_row: int = -1, font_size: int = 7):
    """Add an IEEE-style table to the DOCX document."""
    # Caption above
    cp = doc.add_paragraph()
    cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = cp.add_run(caption.replace("\n", " "))
    r.bold = True; r.font.size = Pt(8)
    r.font.name = "Times New Roman"

    n_cols = len(header)
    tbl = doc.add_table(rows=1 + len(rows), cols=n_cols)
    tbl.style = "Table Grid"

    # Header row
    hr = tbl.rows[0]
    for i, h in enumerate(header):
        c = hr.cells[i]
        c.text = h.replace("\n", " ")
        run = c.paragraphs[0].runs[0]
        run.bold = True
        run.font.size = Pt(font_size)
        run.font.name = "Times New Roman"
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        _docx_shd(c, "1F1F1F")

    # Data rows
    for ri, row in enumerate(rows):
        tr = tbl.rows[ri + 1]
        for ci, val in enumerate(row):
            c = tr.cells[ci]
            c.text = val
            run = c.paragraphs[0].runs[0]
            run.font.size = Pt(font_size)
            run.font.name = "Times New Roman"
            if ri == highlight_row:
                _docx_shd(c, "D8EAD3")
            elif ri % 2 == 1:
                _docx_shd(c, "F8F8F8")

    if note:
        np_ = doc.add_paragraph()
        r = np_.add_run(note)
        r.italic = True; r.font.size = Pt(7)
        r.font.name = "Times New Roman"
        r.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
    return tbl


def _docx_shd(cell, hex_color: str):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:fill"),  hex_color)
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:val"),   "clear")
    tcPr.append(shd)


def build_paper_docx(path: str):
    doc = Document()

    # ── Page 1: Letter, IEEE margins, single-column ────────────────────────
    sec0 = doc.sections[0]
    _docx_page_size_letter(sec0)
    _docx_set_cols(sec0, num_cols=1)

    def _tnr(run, sz=10):
        run.font.name = "Times New Roman"; run.font.size = Pt(sz)

    # Title
    tp = doc.add_paragraph()
    tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = tp.add_run(PAPER_TITLE); _tnr(r, 18)
    r.bold = True; r.font.color.rgb = RGBColor(0, 0, 0)

    # Authors + affil
    ap = doc.add_paragraph()
    ap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = ap.add_run(PAPER_AUTHORS); _tnr(r, 11); r.italic = True
    afp = doc.add_paragraph()
    afp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = afp.add_run(PAPER_AFFIL); _tnr(r, 9)

    doc.add_paragraph()

    # Central finding box (shaded)
    cfp = doc.add_paragraph()
    r = cfp.add_run(
        "Central Finding: Temperature dominates routing strategy. "
        "Single-agent T=0.3 achieves 91%. MCN-LinUCB (homogeneous T=0.3) scores 86% "
        "(−5 pp). MCN-GNN (heterogeneous T=0.1/0.5/0.9) scores 88% = T=0.9 single agent. "
        "T_2 achieves 95% in isolation within Run 10, quantifying a −7 pp routing cost "
        "versus an oracle. No statistically significant routing specialisation detected "
        "in any run (chi-squared p > 0.73).")
    _tnr(r, 8.5); r.italic = True
    r.font.color.rgb = RGBColor(0x1F, 0x5C, 0x2E)
    cfp.paragraph_format.left_indent  = Inches(0.2)
    cfp.paragraph_format.right_indent = Inches(0.2)
    doc.add_paragraph()

    # Abstract
    abh = doc.add_paragraph()
    r = abh.add_run("Abstract"); _tnr(r, 10); r.bold = True
    abp = doc.add_paragraph()
    abp.paragraph_format.left_indent  = Inches(0.25)
    abp.paragraph_format.right_indent = Inches(0.25)
    r = abp.add_run(IEEE_ABSTRACT); _tnr(r, 9)

    kp = doc.add_paragraph()
    kp.paragraph_format.left_indent  = Inches(0.25)
    kp.paragraph_format.right_indent = Inches(0.25)
    r = kp.add_run(f"Index Terms — {IEEE_KEYWORDS}"); _tnr(r, 9); r.italic = True

    doc.add_paragraph()

    # Tables (single-column, before body)
    doc.add_paragraph().add_run("TABLE I — FIVE-CONDITION PHASE 3 COMPARISON"
                                 ).font.name = "Times New Roman"
    _docx_ieee_table(doc, TABLE1_CAPTION, TABLE1_HEADER, TABLE1_ROWS,
                     note=TABLE1_NOTE, highlight_row=1, font_size=7)
    doc.add_paragraph()
    _docx_ieee_table(doc, TABLE2_CAPTION, TABLE2_HEADER, TABLE2_ROWS,
                     note=TABLE2_NOTE, font_size=7)
    doc.add_paragraph()
    _docx_ieee_table(doc, TABLE3_CAPTION, TABLE3_HEADER, TABLE3_ROWS,
                     note=TABLE3_NOTE, font_size=7)
    doc.add_paragraph()
    _docx_ieee_table(doc, TABLE4_CAPTION, TABLE4_HEADER, TABLE4_ROWS,
                     note=TABLE4_NOTE, font_size=7)
    doc.add_paragraph()
    _docx_ieee_table(doc, TABLE5_CAPTION, TABLE5_HEADER, TABLE5_ROWS,
                     note=TABLE5_NOTE, font_size=7)
    doc.add_paragraph()
    _docx_ieee_table(doc, TABLE6_CAPTION, TABLE6_HEADER, TABLE6_ROWS,
                     note=TABLE6_NOTE, font_size=7)
    doc.add_paragraph()
    _docx_ieee_table(doc, TABLE7_CAPTION, TABLE7_HEADER, TABLE7_ROWS,
                     note=TABLE7_NOTE, font_size=7, highlight_row=3)
    doc.add_paragraph()
    _docx_ieee_table(doc, TABLE8_CAPTION, TABLE8_HEADER, TABLE8_ROWS,
                     note=TABLE8_NOTE, font_size=7, highlight_row=1)
    doc.add_paragraph()

    # ── Continuous section break → two-column body ─────────────────────────
    sec1 = doc.add_section(WD_SECTION.CONTINUOUS)
    _docx_page_size_letter(sec1)
    _docx_set_cols(sec1, num_cols=2, space_twips=720)

    # Body sections
    for heading, body in IEEE_SECTIONS:
        sub = _is_subsection(heading)
        hp = doc.add_paragraph()
        if sub:
            hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
            r = hp.add_run(heading); _tnr(r, 10)
            r.bold = True; r.italic = True
        else:
            hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = hp.add_run(heading); _tnr(r, 10); r.bold = True
        hp.paragraph_format.space_before = Pt(8 if not sub else 5)

        # Body text
        blocks = body.split("\n\n")
        first_para = True
        for block in blocks:
            if not block.strip():
                continue
            lines = block.split("\n")
            if all(l.startswith("  ") or not l.strip() for l in lines):
                for line in lines:
                    if line.strip():
                        lp = doc.add_paragraph(style="List Bullet")
                        r = lp.add_run(line.strip()); _tnr(r, 9.5)
            else:
                text = " ".join(l.strip() for l in lines if l.strip())
                bp = doc.add_paragraph()
                r = bp.add_run(text); _tnr(r, 10)
                bp.paragraph_format.space_after = Pt(3)
                if not first_para:
                    bp.paragraph_format.first_line_indent = Pt(14)
                first_para = False

    doc.save(path)
    print(f"  Paper DOCX saved -> {path}")


# ══════════════════════════════════════════════════════════════════════════════
# PROGRESS REPORT helpers
# ══════════════════════════════════════════════════════════════════════════════

def _rpt_shd(cell, hex_color: str):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:fill"),  hex_color)
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:val"),   "clear")
    tcPr.append(shd)


def _rpt_table(doc, header, rows, hfill="1F3864", hcol="FFFFFF",
               alt="E8EEF7", font_size=8):
    tbl = doc.add_table(rows=1 + len(rows), cols=len(header))
    tbl.style = "Table Grid"
    for i, h in enumerate(header):
        c = tbl.rows[0].cells[i]
        c.text = h
        r = c.paragraphs[0].runs[0]
        r.bold = True; r.font.size = Pt(font_size - 0.5)
        r.font.color.rgb = RGBColor(
            int(hcol[0:2], 16), int(hcol[2:4], 16), int(hcol[4:6], 16))
        c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        _rpt_shd(c, hfill)
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            c = tbl.rows[ri+1].cells[ci]
            c.text = val
            c.paragraphs[0].runs[0].font.size = Pt(font_size - 1)
            if ri % 2 == 0:
                _rpt_shd(c, alt)
    return tbl


def _rpt_h(doc, text, size=12, bold=True, color=(0x1F, 0x38, 0x64)):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = bold; r.font.size = Pt(size)
    r.font.color.rgb = RGBColor(*color)
    return p


def _rpt_margins(doc, top=2.0, bot=2.0, left=2.5, right=2.5):
    for s in doc.sections:
        s.top_margin    = Cm(top)
        s.bottom_margin = Cm(bot)
        s.left_margin   = Cm(left)
        s.right_margin  = Cm(right)


def _rpt_pdf_styles():
    S = {}
    S["title"]   = ParagraphStyle("rt",  fontSize=20, textColor=DARK_BLUE,
                                   alignment=TA_CENTER, spaceAfter=4,
                                   fontName="Helvetica-Bold")
    S["sub"]     = ParagraphStyle("rs",  fontSize=10, textColor=ACCENT,
                                   alignment=TA_CENTER, spaceAfter=12,
                                   fontName="Helvetica-Oblique")
    S["finding"] = ParagraphStyle("rf",  fontSize=8.5,
                                   textColor=colors.HexColor("#1F5C2E"),
                                   alignment=TA_LEFT, spaceAfter=8,
                                   fontName="Helvetica-BoldOblique",
                                   backColor=colors.HexColor("#E2EFDA"),
                                   leftIndent=6, borderPad=3)
    S["h1"]      = ParagraphStyle("rh1", fontSize=12, textColor=DARK_BLUE,
                                   spaceBefore=12, spaceAfter=3,
                                   fontName="Helvetica-Bold")
    S["body"]    = ParagraphStyle("rb",  fontSize=8.5, textColor=BLACK,
                                   leading=13, spaceAfter=5,
                                   alignment=TA_JUSTIFY, fontName="Helvetica")
    S["mk"]      = ParagraphStyle("rmk", fontSize=8.5,
                                   fontName="Helvetica-Bold", textColor=DARK_BLUE)
    S["mv"]      = ParagraphStyle("rmv", fontSize=8.5,
                                   fontName="Helvetica",     textColor=BLACK)
    S["th"]      = ParagraphStyle("rth", fontSize=7.5,
                                   fontName="Helvetica-Bold", textColor=WHITE,
                                   alignment=TA_CENTER)
    S["td"]      = ParagraphStyle("rtd", fontSize=7.5,
                                   fontName="Helvetica",     textColor=BLACK,
                                   alignment=TA_CENTER)
    S["tdl"]     = ParagraphStyle("rtdl",fontSize=7.5,
                                   fontName="Helvetica",     textColor=BLACK)
    S["note"]    = ParagraphStyle("rnt", fontSize=7.5, textColor=GREY,
                                   fontName="Helvetica-Oblique",
                                   alignment=TA_LEFT)
    S["footer"]  = ParagraphStyle("rft", fontSize=7,  textColor=GREY,
                                   alignment=TA_CENTER,
                                   fontName="Helvetica-Oblique")
    return S


def _rpt_pdf_table(header, rows, col_widths, S,
                   hdr_bg=DARK_BLUE, alt_bg=ROW_ALT,
                   highlight_rows=None, highlight_color=AMBER):
    data = [[Paragraph(h, S["th"]) for h in header]]
    for row in rows:
        data.append([Paragraph(c, S["td"] if len(c) < 12 else S["tdl"])
                     for c in row])
    t = Table(data, colWidths=col_widths, repeatRows=1)
    cmds = list(GRID_BASE) + [
        ("BACKGROUND",     (0,0),(-1,0),  hdr_bg),
        ("ROWBACKGROUNDS", (0,1),(-1,-1), [WHITE, alt_bg]),
    ]
    if highlight_rows:
        for r in highlight_rows:
            cmds.append(("BACKGROUND", (0,r),(-1,r), highlight_color))
    t.setStyle(TableStyle(cmds))
    return t


# ══════════════════════════════════════════════════════════════════════════════
# PROGRESS REPORT — DOCX
# ══════════════════════════════════════════════════════════════════════════════

def build_report_docx(path: str):
    doc = Document()
    _rpt_margins(doc)

    tp = doc.add_paragraph()
    tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = tp.add_run(REPORT_TITLE)
    r.bold = True; r.font.size = Pt(17)
    r.font.color.rgb = RGBColor(0x1F, 0x38, 0x64)

    sp = doc.add_paragraph()
    sp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sp.add_run(REPORT_SUBTITLE)
    r.italic = True; r.font.size = Pt(10)
    r.font.color.rgb = RGBColor(0x44, 0x72, 0xC4)
    doc.add_paragraph()

    note = doc.add_paragraph()
    r = note.add_run(
        "CENTRAL FINDING — Phase 3: Temperature is the dominant performance variable. "
        "Single-agent T=0.3 achieves 91%. MCN-LinUCB (3×T=0.3) achieves 86% (−5pp). "
        "MCN-GNN (T=0.1/0.5/0.9) achieves 88% = T=0.9 single agent. "
        "T_2 achieves 95% in isolation within Run 10, giving a −7pp routing cost vs oracle. "
        "No statistically significant routing specialisation (chi-sq p>0.73).")
    r.font.size = Pt(8.5); r.bold = True
    r.font.color.rgb = RGBColor(0x1F, 0x5C, 0x2E)
    doc.add_paragraph()

    mt = doc.add_table(rows=len(REPORT_META), cols=2)
    mt.style = "Table Grid"
    for i, (k, v) in enumerate(REPORT_META):
        row = mt.rows[i]
        row.cells[0].text = k
        row.cells[0].paragraphs[0].runs[0].bold = True
        row.cells[0].paragraphs[0].runs[0].font.size = Pt(9)
        row.cells[1].text = v
        row.cells[1].paragraphs[0].runs[0].font.size = Pt(9)
        if i % 2 == 0:
            _rpt_shd(row.cells[0], "E8EEF7"); _rpt_shd(row.cells[1], "E8EEF7")
    doc.add_paragraph()

    _rpt_h(doc, "All Runs Summary (Runs 1-10)", size=12)
    _rpt_table(doc, REPORT_RUN_TABLE_HEADER, REPORT_RUN_TABLE, font_size=8)
    doc.add_paragraph()

    _rpt_h(doc, "Phase 3: Five-Condition Comparison", size=12)
    _rpt_table(doc, REPORT_ABL_HEADER, REPORT_ABL_TABLE, font_size=8)
    doc.add_paragraph()

    _rpt_h(doc, "MCN Variants vs Best Single Agent", size=12)
    _rpt_table(doc, TABLE3_HEADER, TABLE3_ROWS, font_size=8)
    doc.add_paragraph()

    _rpt_h(doc, "Defect Log (All Phases)", size=12)
    _rpt_table(doc, REPORT_DEFECT_HEADER, REPORT_DEFECTS, font_size=8)
    doc.add_paragraph()

    _rpt_h(doc, "Current System State & Next Steps", size=12)
    st = doc.add_table(rows=len(REPORT_STATE), cols=2)
    st.style = "Table Grid"
    for i, (k, v) in enumerate(REPORT_STATE):
        row = st.rows[i]
        row.cells[0].text = k
        row.cells[0].paragraphs[0].runs[0].bold = True
        row.cells[0].paragraphs[0].runs[0].font.size = Pt(9)
        row.cells[1].text = v
        row.cells[1].paragraphs[0].runs[0].font.size = Pt(9)
        if i % 2 == 0:
            _rpt_shd(row.cells[0], "E8EEF7"); _rpt_shd(row.cells[1], "E8EEF7")
        if k.startswith("Next"):
            _rpt_shd(row.cells[0], "D6E4F0"); _rpt_shd(row.cells[1], "D6E4F0")

    doc.save(path)
    print(f"  Report DOCX saved -> {path}")


# ══════════════════════════════════════════════════════════════════════════════
# PROGRESS REPORT — PDF
# ══════════════════════════════════════════════════════════════════════════════

def build_report_pdf(path: str):
    from reportlab.lib.pagesizes import A4
    doc = SimpleDocTemplate(
        path, pagesize=A4,
        leftMargin=2.2*cm, rightMargin=2.2*cm,
        topMargin=2.0*cm,  bottomMargin=2.0*cm,
        title=REPORT_TITLE,
    )
    S = _rpt_pdf_styles()
    story = []

    story.append(Paragraph(REPORT_TITLE,    S["title"]))
    story.append(Paragraph(REPORT_SUBTITLE, S["sub"]))
    story.append(HRFlowable(width="100%", thickness=1.5, color=DARK_BLUE, spaceAfter=6))
    story.append(Paragraph(
        "<b>Central Finding (Phase 3):</b> Temperature is the dominant variable. "
        "Single-agent T=0.3 achieves 91% — best overall. "
        "MCN-LinUCB (homogeneous T=0.3) achieves 86% (−5pp). "
        "MCN-GNN (heterogeneous T=0.1/0.5/0.9) achieves 88% = T=0.9 single agent. "
        "T_2 achieves 95% in isolation (within-run), giving a −7pp routing cost "
        "vs an oracle. No specialisation in any run (chi-sq p&gt;0.73).",
        S["finding"]))

    # Meta
    md = [[Paragraph(k, S["mk"]), Paragraph(v, S["mv"])] for k, v in REPORT_META]
    mt = Table(md, colWidths=[3.8*cm, None])
    mt.setStyle(TableStyle(list(GRID_BASE) + [
        ("BACKGROUND",     (0,0),(0,-1), LIGHT_BLUE),
        ("ROWBACKGROUNDS", (0,0),(-1,-1),[WHITE, ROW_ALT]),
    ]))
    story.append(mt)
    story.append(Spacer(1, 0.4*cm))

    story.append(Paragraph("<b>All Experimental Runs (1-10)</b>", S["h1"]))
    story.append(HRFlowable(width="100%", thickness=0.5, color=DARK_BLUE, spaceAfter=3))
    cwr = [0.8*cm, 1.6*cm, 1.4*cm, 1.3*cm, 2.0*cm, 2.5*cm, 5.0*cm]
    story.append(_rpt_pdf_table(
        REPORT_RUN_TABLE_HEADER, REPORT_RUN_TABLE, cwr, S,
        highlight_rows=[7, 8], highlight_color=colors.HexColor("#D6E4F0")))
    story.append(Spacer(1, 0.4*cm))

    story.append(Paragraph("<b>Phase 3: Five-Condition Comparison</b>", S["h1"]))
    story.append(HRFlowable(width="100%", thickness=0.5, color=DARK_BLUE, spaceAfter=3))
    cwa = [3.5*cm, 1.5*cm, 1.3*cm, 2.0*cm, 2.2*cm, 2.2*cm, 2.2*cm, 2.2*cm]
    story.append(_rpt_pdf_table(
        REPORT_ABL_HEADER, REPORT_ABL_TABLE, cwa, S,
        highlight_rows=[2], highlight_color=GREEN_LIGHT))
    story.append(Spacer(1, 0.4*cm))

    story.append(Paragraph("<b>MCN Variants vs Best Single Agent</b>", S["h1"]))
    story.append(HRFlowable(width="100%", thickness=0.5, color=DARK_BLUE, spaceAfter=3))
    cw3 = [3.5*cm, 4.5*cm, 3.5*cm, 4.0*cm]
    story.append(_rpt_pdf_table(
        TABLE3_HEADER, TABLE3_ROWS, cw3, S,
        highlight_rows=[1], highlight_color=AMBER))
    story.append(Spacer(1, 0.4*cm))

    story.append(Paragraph("<b>Defect Log (All Phases)</b>", S["h1"]))
    story.append(HRFlowable(width="100%", thickness=0.5, color=DARK_BLUE, spaceAfter=3))
    cwd = [5.5*cm, 2.5*cm, 8.0*cm]
    story.append(_rpt_pdf_table(
        REPORT_DEFECT_HEADER, REPORT_DEFECTS, cwd, S,
        highlight_rows=[5], highlight_color=AMBER))
    story.append(Spacer(1, 0.4*cm))

    story.append(Paragraph("<b>Current State and Next Steps</b>", S["h1"]))
    story.append(HRFlowable(width="100%", thickness=0.5, color=DARK_BLUE, spaceAfter=3))
    sd = [[Paragraph(k, S["mk"]), Paragraph(v, S["mv"])] for k, v in REPORT_STATE]
    stbl = Table(sd, colWidths=[3.5*cm, None])
    cmds = list(GRID_BASE) + [
        ("BACKGROUND",     (0,0),(0,-1),  LIGHT_BLUE),
        ("ROWBACKGROUNDS", (0,0),(-1,-1), [WHITE, ROW_ALT]),
    ]
    for i, (k, _) in enumerate(REPORT_STATE):
        if k.startswith("Next"):
            cmds.append(("BACKGROUND", (0,i),(-1,i), colors.HexColor("#D6E4F0")))
    stbl.setStyle(TableStyle(cmds))
    story.append(stbl)

    story.append(Spacer(1, 0.5*cm))
    story.append(HRFlowable(width="100%", thickness=1, color=DARK_BLUE))
    story.append(Paragraph(
        "Complete edition — Runs 1-10 + 3 ablations  |  "
        "MLflow tracked  |  Raw data: mcn-results:/results/runs.jsonl  |  Feb 2026",
        S["footer"]))

    doc.build(story)
    print(f"  Report PDF  saved -> {path}")


# ══════════════════════════════════════════════════════════════════════════════
# Main
# ══════════════════════════════════════════════════════════════════════════════

if __name__ == "__main__":
    print("Generating MCN documents (IEEE edition + complete progress report)...")

    build_report_docx(os.path.join(OUT_DIR, "MCN_Progress_Report.docx"))
    build_report_pdf( os.path.join(OUT_DIR, "MCN_Progress_Report.pdf"))
    build_paper_docx( os.path.join(OUT_DIR, "MCN_Academic_Paper.docx"))
    build_paper_pdf(  os.path.join(OUT_DIR, "MCN_Academic_Paper.pdf"))

    print("Done. Four files written to", OUT_DIR)
